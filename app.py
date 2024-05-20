from flask import Flask, render_template, request, redirect, url_for, session, flash, make_response
import sqlite3
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import os

app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Database connection function
def get_db_connection():
    database_path = os.path.join('/var/data', 'shifts.db')
    conn = sqlite3.connect(database_path)
    conn.row_factory = sqlite3.Row
    return conn

# User authentication middleware
def login_required(func):
    def wrapper(*args, **kwargs):
        if 'username' in session:
            return func(*args, **kwargs)
        else:
            return redirect(url_for('login'))
    return wrapper

# Main route to submit shift details
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        branch = request.form['branch']
        staff_name = branch
        staff_number = branch + '123'
        mobile = request.form['mobile']
        shift_timing = request.form['shift_timing']

        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO shifts (branch, staff_name, staff_number, mobile, shift_timing) 
            VALUES (?, ?, ?, ?, ?)
        ''', (branch, staff_name, staff_number, mobile, shift_timing))
        conn.commit()
        conn.close()

        flash('Shift details submitted successfully', 'success')
        return redirect(url_for('index'))
    else:
        return render_template('index.html')

# Admin dashboard route
@app.route('/admin')
@login_required
def admin_dashboard():
    if session['username'] == 'admin':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM shifts')
        shifts = cursor.fetchall()
        conn.close()
        return render_template('admin_dashboard.html', shifts=shifts)
    else:
        flash('Access denied! You do not have permission to view this page.', 'error')
        return redirect(url_for('index'))

# Login route
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        if username == 'admin' and password == 'admin':
            session['username'] = 'admin'
            flash('Logged in successfully!', 'success')
            return redirect(url_for('admin_dashboard'))
        else:
            branch = username
            branch_password = branch + '123'
            if password == branch_password:
                session['username'] = branch
                flash('Logged in successfully!', 'success')
                return redirect(url_for('index'))
            else:
                flash('Invalid username or password', 'error')
                return redirect(url_for('login'))

    return render_template('login.html')

# Logout route
@app.route('/logout')
def logout():
    session.pop('username', None)
    flash('Logged out successfully!', 'success')
    return redirect(url_for('login'))

# Generate PDF report
@app.route('/generate_pdf')
@login_required
def generate_pdf():
    if session['username'] == 'admin':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM shifts')
        shifts = cursor.fetchall()
        conn.close()

        response = make_response('')
        response.headers['Content-Disposition'] = 'attachment; filename=shifts_report.pdf'

        pdf = canvas.Canvas(response)
        pdf.setFont("Helvetica", 12)
        pdf.drawString(100, 750, "Shifts Report")

        y = 700
        for shift in shifts:
            y -= 20
            pdf.drawString(100, y, f"Branch: {shift['branch']}, Staff Name: {shift['staff_name']}, Staff Number: {shift['staff_number']}, Mobile: {shift['mobile']}, Shift Timing: {shift['shift_timing']}")

        pdf.save()
        return response
    else:
        flash('Access denied! You do not have permission to view this page.', 'error')
        return redirect(url_for('index'))

# Generate DOCX report
@app.route('/generate_docx')
@login_required
def generate_docx():
    if session['username'] == 'admin':
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM shifts')
        shifts = cursor.fetchall()
        conn.close()

        document = Document()
        document.add_heading('Shifts Report', level=1)

        for shift in shifts:
            document.add_paragraph(f"Branch: {shift['branch']}, Staff Name: {shift['staff_name']}, Staff Number: {shift['staff_number']}, Mobile: {shift['mobile']}, Shift Timing: {shift['shift_timing']}")

        response = make_response('')
        response.headers['Content-Disposition'] = 'attachment; filename=shifts_report.docx'
        document.save(response)
        return response
    else:
        flash('Access denied! You do not have permission to view this page.', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)
